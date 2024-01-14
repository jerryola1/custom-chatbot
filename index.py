import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set document title
title = doc.add_heading('Curriculum Vitae', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add your personal information
doc.add_heading('Personal Information', level=1)
personal_info = doc.add_paragraph()
personal_info.add_run('Name: ').bold = True
personal_info.add_run('Your Name').bold = False
personal_info.add_run('\n')
personal_info.add_run('Address: ').bold = True
personal_info.add_run('Your Address').bold = False
personal_info.add_run('\n')
personal_info.add_run('Phone: ').bold = True
personal_info.add_run('Your Phone Number').bold = False
personal_info.add_run('\n')
personal_info.add_run('Email: ').bold = True
personal_info.add_run('Your Email Address').bold = False
personal_info.add_run('\n')

# Add your education section
doc.add_heading('Education', level=1)
edu_section = doc.add_paragraph()
edu_section.add_run('Degree: ').bold = True
edu_section.add_run('Your Degree').bold = False
edu_section.add_run('\n')
edu_section.add_run('Institution: ').bold = True
edu_section.add_run('University/College Name').bold = False
edu_section.add_run('\n')
edu_section.add_run('Year: ').bold = True
edu_section.add_run('Year of Graduation').bold = False
edu_section.add_run('\n')

# Add your work experience section
doc.add_heading('Work Experience', level=1)
work_section = doc.add_paragraph()
work_section.add_run('Job Title: ').bold = True
work_section.add_run('Your Job Title').bold = False
work_section.add_run('\n')
work_section.add_run('Company: ').bold = True
work_section.add_run('Company Name').bold = False
work_section.add_run('\n')
work_section.add_run('Year: ').bold = True
work_section.add_run('Year of Employment').bold = False
work_section.add_run('\n')

# Define the file name
file_name = 'cv_uk_format.docx'

# Save the document
doc.save(file_name)

# Open the document for viewing using the default application
os.system(file_name)

print('CV created and opened for viewing successfully!')
